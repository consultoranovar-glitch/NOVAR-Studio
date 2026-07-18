#!/usr/bin/env python3
"""
Motor de documentos NOVAR · Markdown → Word con estilo editorial NOVAR.

Uso: python3 md2docx.py entrada.md salida.docx [--tokens marcas/novar/design-tokens.json]

Soporta: # ## ### títulos · **negrita** · *itálica* · tablas | · listas - y 1. · > citas.
Tipografías: títulos Georgia, cuerpo Arial (estándar editorial NOVAR probado).
SIEMPRE validar la salida con validar_ooxml.py antes de entregar.
"""
import argparse, json, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches

def rgb(hx): return RGBColor(int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16))

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("md"); ap.add_argument("docx")
    ap.add_argument("--tokens", default=None)
    a = ap.parse_args()
    C = {"gold": "C9A84C", "goldDark": "A07F33", "ink": "1A1A1A", "navy": "0A1A66", "gray": "6E6E6E"}
    if a.tokens:
        C.update(json.load(open(a.tokens, encoding="utf-8")).get("colores", {}))
    GOLDD, INK, NAVY, GRAY = rgb(C["goldDark"]), rgb(C["ink"]), rgb(C["navy"]), rgb(C["gray"])

    doc = Document()
    sc = doc.sections[0]
    sc.top_margin = Inches(1.0); sc.bottom_margin = Inches(0.9)
    sc.left_margin = Inches(1.05); sc.right_margin = Inches(1.05)

    def add_runs(p, text, base_italic=False):
        for part in re.split(r'(\*\*.*?\*\*|\*.*?\*)', text):
            if not part: continue
            r = p.add_run()
            if part.startswith("**") and part.endswith("**"): r.text = part[2:-2]; r.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) > 2: r.text = part[1:-1]; r.italic = True
            else: r.text = part
            if base_italic: r.italic = True
            r.font.name = "Arial"; r.font.size = Pt(10.5); r.font.color.rgb = INK

    lines = open(a.md, encoding="utf-8").read().splitlines()
    i = 0
    while i < len(lines):
        l = lines[i].rstrip()
        if l.startswith("# "):
            p = doc.add_paragraph(); r = p.add_run(l[2:]); r.bold = True
            r.font.name = "Georgia"; r.font.size = Pt(22); r.font.color.rgb = INK
        elif l.startswith("## "):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16)
            r = p.add_run(l[3:]); r.bold = True
            r.font.name = "Georgia"; r.font.size = Pt(15); r.font.color.rgb = NAVY
        elif l.startswith("### "):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
            r = p.add_run(l[4:].replace("*", "")); r.bold = True
            r.font.name = "Georgia"; r.font.size = Pt(12.5); r.font.color.rgb = GOLDD
        elif l.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(set(c) <= set("-: ") for c in cells): rows.append(cells)
                i += 1
            i -= 1
            t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"
            for ri, rw in enumerate(rows):
                for ci, cv in enumerate(rw):
                    cp = t.cell(ri, ci).paragraphs[0]; add_runs(cp, cv)
                    if ri == 0:
                        for r in cp.runs: r.bold = True
        elif l.startswith("> "):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
            add_runs(p, l[2:], base_italic=True)
            for r in p.runs: r.font.color.rgb = GRAY
        elif l.startswith("- ") or re.match(r"^\d+\. ", l):
            txt = re.sub(r"^(\- |\d+\. )", "", l)
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(14)
            pre = p.add_run("•  " if l.startswith("- ") else l.split(" ")[0] + "  ")
            pre.bold = True; pre.font.name = "Arial"; pre.font.size = Pt(10.5); pre.font.color.rgb = GOLDD
            add_runs(p, txt)
        elif l == "---" or not l.strip():
            pass
        else:
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.25
            add_runs(p, l.replace("`", ""))
        i += 1
    doc.save(a.docx)
    print("OK", a.docx, "— valida con validar_ooxml.py antes de entregar")

if __name__ == "__main__":
    main()
